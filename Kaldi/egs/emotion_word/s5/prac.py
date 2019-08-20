#!/usr/bin/env python        
# -*- coding: utf-8 -*-    
import numpy as np
import scipy.io as sio
import scipy.io.wavfile
import matplotlib.pyplot as plt
import librosa
import librosa.display
from PIL import Image
import PIL.Image as pilimg
import docx
from docx.oxml.ns import qn, nsdecls
import scipy.io.wavfile as wav
from speechpy.feature import mfcc
from docx.shared import RGBColor


import tensorflow as tf
import keras
from keras.models import Sequential
from keras.layers import Dense, Dropout, Flatten
from keras.layers.convolutional import Conv2D, MaxPooling2D
from keras.models import load_model, Model

def get_feature_vector_from_mfcc(file_path: str, flatten: bool,
                                 mfcc_len: int = 39) -> np.ndarray:
    fs, signal = wav.read(file_path)
    #print(fs)
    #print(np.shape(signal))
    #signal=down_sample(file_path, fs, 8000)
    #fs=8000
    s_len = len(signal)
    # pad the signals to have same size if lesser than required
    # else slice them
    if s_len < mean_signal_length:
        pad_len = mean_signal_length - s_len
        pad_rem = pad_len % 2
        pad_len //= 2
        signal = np.pad(signal, (pad_len, pad_len + pad_rem),
                        'constant', constant_values=0)
    else:
        pad_len = s_len - mean_signal_length
        pad_len //= 2
        signal = signal[pad_len:pad_len + mean_signal_length]
    mel_coefficients = mfcc(signal, fs, num_cepstral=mfcc_len)
    #if flatten:
        # Flatten the data
    #    mel_coefficients = np.ravel(mel_coefficients)
    return mel_coefficients


### 문자출력 + 문자수 출력  ###
L=[]
text=[]
file=open('./a.txt','r')
while (1):
    line=file.readline()

    try:escape=line.index('\n')
    except:escape=len(line)

    if line:
        L.append(line[0:escape])
    else:
        break

file.close()
text=L[-1][53:]
text=text.strip()
print("음성인식 결과 : {0}" . format( text ))
print("글자 수 : %d " . format(len(text)))



jump=[]
for i in range(0, len(text)):
    if text[i]==" ":
         
         jump.append(i)
spacing=[]
if len(jump) >0:
    for i in range (0, len(jump)):
        if i == 0:
            spacing.append(text[0: int(jump[0])])
        else:
            spacing.append(text[ jump[i-1]+1 :jump[i]]) 
    spacing.append(text[jump[i-1]+1:])

num_spacing=len(spacing)
for i in range(0, len(spacing)):
    print("{0} 번째 단어 : {1}" . format((i+1), spacing[i]))
   
print(spacing)
print(len(spacing))

samplerate, data = sio.wavfile.read('./01r6020005Child_0001.wav')

data_s=data # 기본 파장 
plt.subplot(6,1,1)
plt.plot(data_s)


data_abs=np.log(abs(data)) # 파장 절대값

def relu(x):
    if x <= 8:
        return 0
    return x

a=[]
for i in range(0, len(data)):
    a.append(relu(data_abs[i]))

data=a
data_del8=a # 8 밑으로 지운 것
plt.subplot(6,1,2)
plt.plot(data_del8)


data_main=[]
for i in range(0, num_spacing):
    data_main.append([i])


standard=1000
a=[]
if num_spacing %2 == 0:
    for i in range (0, int(num_spacing/2)):
        #front
        front=[]
        k=0
        while len(front) <1:
            if data_del8[k] !=0:
                front.append(k)
            k += 1

        back=[]
        k=0
        while len(back) <1:
            if data_del8[-k] !=0:
                back.append(len(data_del8)-k)
            k += 1

        data_slice=front+back
        data_s=data_del8[data_slice[0]:data_slice[1]]
        
        data_section=[]
        for n in range(0, len(data_s)-standard):
            if data_s[n:standard+n].count(0)==standard:
                  data_section.append(n)
                 # print(data_section)
                 # print(data_s[0:data_section[0]]) 
        #print(data_section)
        data_main[i]=data_s[0:data_section[0]]
        data_main[num_spacing-1-i]=data_s[data_section[-1]+standard:]
        
        data_del8=data_s[data_section[0]:data_section[-1]+standard:]

if num_spacing %2 !=0 :
    for i in range (0, int(num_spacing/2)):
        #front
        front=[]
        k=0
        while len(front) <1:
            if data_del8[k] !=0:
                front.append(k)
            k += 1

        back=[]
        k=0
        while len(back) <1:
            if data_del8[-k] !=0:
                back.append(len(data_del8)-k)
            k += 1

        data_slice=front+back
        data_s=data_del8[data_slice[0]:data_slice[1]]

        if i != num_spacing:
            data_section=[]
            for n in range(0, len(data_s)-standard):
                if data_s[n:standard+n].count(0)==standard:
                    data_section.append(n)
            
            data_main[i]=data_s[0:data_section[0]]
            data_main[num_spacing-1-i]=data_s[data_section[-1]+standard:]
            data_del8=data_s[data_section[0]:data_section[-1]+standard:]    
        elif i == num_spacing:
            data_main[i]=data_s 

    



for n in range(0, len(data_main)):

    a=[]
    for i in range(0, len(data_main[n])):
        if data_main[n][i] != 0:
            a.append(data_main[n][i])
    data_main[n]=a



seperation=[]
for i in range(0, num_spacing):
    a=[]
    for n in range(0, len(spacing[i])):
        k=int(len(data_main[i])/len(spacing[i]))
        a.append(sum(data_main[0][n*k:(n+1)*(k)])/k)
    seperation.append(a)



mean_signal_length=32000
to_flatten=True
signal=get_feature_vector_from_mfcc('./01r6020005Child_0001.wav', flatten=to_flatten)

img_rows = 398
img_cols = 39
#input_shape = (img_rows, img_cols, 1)

signal = signal.reshape(1, img_rows, img_cols)
print(np.shape(signal))

model = load_model('./emotion_LSTM.h5')
model.summary()
emotion=model.predict_classes(signal)
print("0:Angry, 1:Happy, 2:Neutral, 3:Sad, 4:Disgust, 5:Lasy, 6:Scared") 
print(emotion)



if emotion == 0:
    font_name = "이순신 Regular"
    expression = "ㅡㅡ"

if emotion == 1:
    font_name = "은 필기"
    expression = "ㅋㅋㅋㅋㅋ"

if emotion == 2:
    font_name = "이순신 돋움체 M"
    expression = "."

if emotion == 3:
    font_name = "산돌종이학Blur"
    expression = "...ㅠㅠ"

if emotion == 4:
    font_name = "HY궁서"
    expression = "ㅡ.ㅡ"

if emotion == 5:
    font_name = "은 궁서"
    expression = "...."

if emotion == 6:
    font_name = "-흔적L"
    expression = "!?!?!?!?!"


doc= docx.Document()
para = doc.add_paragraph()

print(spacing)
for i in range(0, num_spacing):
    run = para.add_run(" ")
    for n in range(0, len(spacing[i])):
        run = para.add_run(spacing[i][n])
        run.font.size = docx.shared.Pt(10*(int(seperation[i][n])-7))
        #run.bold = True
        run.font.name = font_name
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        #run.font.color.rgb = RGBColor(255, 0, 0)

run = para.add_run(expression)
run.font.name = font_name
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), font_name)

doc.save('./emotiontext.docx')

